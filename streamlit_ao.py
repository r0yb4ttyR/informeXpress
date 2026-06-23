import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from io import BytesIO
from datetime import datetime, date
import json
import base64

st.set_page_config(
    page_title="Generador AO",
    page_icon="🕵️",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={"About": "Generador de Informes de Análisis Operativo (AO)"},
)

# Meses en español para formatear fechas sin depender del locale del sistema
MESES_ES = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]


def fecha_es(d):
    """Formatea una fecha como 'DD de mes de AAAA' en español."""
    return f"{d.day:02d} de {MESES_ES[d.month - 1]} de {d.year}"


def hora_key(evento):
    """Clave para ordenar eventos cronológicamente (los sin hora van primero)."""
    return evento.get('hora') or datetime.min.time()


# Niveles de clasificación disponibles para el informe
CLASIFICACIONES = ["CONFIDENCIAL", "USO INTERNO", "RESERVADO", "SIN CLASIFICAR"]


def _set_cell_shading(cell, fill_hex):
    """Aplica un color de fondo a una celda de tabla DOCX."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def _add_field(paragraph, instr):
    """Inserta un campo de Word (p. ej. PAGE / NUMPAGES) en el párrafo."""
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), instr)
    paragraph._p.append(fld)


def _apply_table_style(table, *names):
    """Aplica el primer estilo de tabla disponible de la lista dada."""
    for name in names:
        try:
            table.style = name
            return
        except KeyError:
            continue


def maps_url(lat, lon):
    """Devuelve una URL de mapa para unas coordenadas (solo texto, sin llamadas externas)."""
    return f"https://www.google.com/maps?q={lat},{lon}"


def _add_hyperlink(paragraph, url, text, color="1F4E79"):
    """Inserta un hiperenlace externo (subrayado) en un párrafo DOCX."""
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), color)
    rPr.append(color_el)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def build_docx(nombre_ao, fecha_operacion, zona, clasificacion,
               observaciones, eventos, interacciones):
    """Genera el informe en formato DOCX y devuelve los bytes."""
    doc = Document()

    # Propiedades del documento
    core = doc.core_properties
    core.title = f"AO - {nombre_ao}" if nombre_ao else "Análisis Operativo"
    core.author = "Generador AO"
    core.subject = "Informe operativo"
    core.created = datetime.now()

    section = doc.sections[0]

    # Encabezado: banner de clasificación en todas las páginas
    header_p = section.header.paragraphs[0]
    header_p.text = clasificacion
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header_p.runs:
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Pie: "Página X de Y" centrado en todas las páginas
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.add_run("Página ").font.size = Pt(9)
    _add_field(footer_p, " PAGE ")
    footer_p.add_run(" de ").font.size = Pt(9)
    _add_field(footer_p, " NUMPAGES ")

    # Título
    titulo = doc.add_heading(f'AO — "{nombre_ao}"', level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Tabla de metadatos
    tabla_meta = doc.add_table(rows=0, cols=2)
    _apply_table_style(tabla_meta, "Light List Accent 1", "Table Grid")
    for etiqueta, valor in [
        ("Fecha", fecha_es(fecha_operacion)),
        ("Zona de operaciones", zona or "—"),
        ("Clasificación", clasificacion),
        ("Generado", datetime.now().strftime("%d/%m/%Y %H:%M")),
    ]:
        celdas = tabla_meta.add_row().cells
        celdas[0].text = etiqueta
        celdas[1].text = str(valor)
        for par in celdas[0].paragraphs:
            for r in par.runs:
                r.bold = True
        _set_cell_shading(celdas[0], "F2F2F2")

    doc.add_paragraph()

    # Observaciones
    doc.add_heading("Observaciones clave", level=1)
    doc.add_paragraph(observaciones or "Sin observaciones.")

    # Cronología como tabla
    doc.add_heading("Cronología de movimientos", level=1)
    if eventos:
        tabla = doc.add_table(rows=1, cols=2)
        _apply_table_style(tabla, "Table Grid")
        encabezado = tabla.rows[0].cells
        encabezado[0].text = "Hora"
        encabezado[1].text = "Descripción"
        for celda in encabezado:
            for par in celda.paragraphs:
                for r in par.runs:
                    r.bold = True
            _set_cell_shading(celda, "D9D9D9")
        for evento in eventos:
            hora_txt = evento['hora'].strftime('%H:%M') if evento.get('hora') else "—"
            celdas = tabla.add_row().cells
            celdas[0].text = hora_txt
            celdas[1].text = evento.get('descripcion') or ""
            if evento.get('lat') is not None and evento.get('lon') is not None:
                par_loc = celdas[1].add_paragraph()
                par_loc.add_run("📍 ")
                _add_hyperlink(par_loc, maps_url(evento['lat'], evento['lon']),
                               f"{evento['lat']:.5f}, {evento['lon']:.5f}")
            if evento.get('imagen'):
                par_img = celdas[1].add_paragraph()
                par_img.add_run().add_picture(BytesIO(evento['imagen']), width=Inches(2.8))
        tabla.autofit = False
        for fila in tabla.rows:
            fila.cells[0].width = Inches(1.0)
            fila.cells[1].width = Inches(5.5)
    else:
        doc.add_paragraph("Sin eventos registrados.")

    # Interacciones
    doc.add_heading("Interacciones", level=1)
    doc.add_paragraph(interacciones or "Sin interacciones.")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _latin1(text):
    """Sanea texto para las fuentes core de FPDF (codificación latin-1)."""
    if not text:
        return ""
    reemplazos = {
        "–": "-", "—": "-", "“": '"', "”": '"',
        "‘": "'", "’": "'", "…": "...", "•": "-", "\u00a0": " ",
    }
    for k, v in reemplazos.items():
        text = text.replace(k, v)
    return text.encode("latin-1", "replace").decode("latin-1")


def build_pdf(nombre_ao, fecha_operacion, zona, clasificacion,
              observaciones, eventos, interacciones):
    """Genera el informe en formato PDF y devuelve los bytes. Requiere fpdf2."""
    from fpdf import FPDF  # import perezoso: solo necesario al exportar a PDF
    from fpdf.enums import XPos, YPos

    GRIS = (120, 120, 120)
    AZUL = (31, 78, 121)

    class _PDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "B", 8)
            self.set_text_color(*GRIS)
            self.cell(0, 6, _latin1(clasificacion), align="C",
                      new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            self.set_draw_color(*GRIS)
            self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
            self.ln(2)
            self.set_text_color(0, 0, 0)

        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(*GRIS)
            self.cell(0, 10, _latin1(f"Página {self.page_no()} de {{nb}}"), align="C")

    pdf = _PDF()
    pdf.set_margins(left=15, top=15, right=15)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.alias_nb_pages()
    pdf.add_page()

    # Título
    pdf.set_text_color(*AZUL)
    pdf.set_font("Helvetica", "B", 18)
    pdf.multi_cell(0, 10, _latin1(f'AO - "{nombre_ao}"'),
                   new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(2)

    # Metadatos
    for etiqueta, valor in [
        ("Fecha", fecha_es(fecha_operacion)),
        ("Zona de operaciones", zona or "—"),
        ("Clasificación", clasificacion),
        ("Generado", datetime.now().strftime("%d/%m/%Y %H:%M")),
    ]:
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(45, 7, _latin1(f"{etiqueta}:"))
        pdf.set_font("Helvetica", "", 11)
        pdf.multi_cell(0, 7, _latin1(str(valor)),
                       new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(3)

    def _seccion(titulo_seccion):
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_text_color(*AZUL)
        pdf.multi_cell(0, 8, _latin1(titulo_seccion),
                       new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_text_color(0, 0, 0)

    # Observaciones
    _seccion("Observaciones clave")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6, _latin1(observaciones or "Sin observaciones."),
                   new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(3)

    # Cronología
    _seccion("Cronología de movimientos")
    if eventos:
        for evento in eventos:
            hora_txt = evento['hora'].strftime('%H:%M') if evento.get('hora') else "—"
            pdf.set_font("Helvetica", "B", 11)
            pdf.cell(16, 6, _latin1(hora_txt))
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 6, _latin1(evento.get('descripcion') or ""),
                           new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            if evento.get('lat') is not None and evento.get('lon') is not None:
                pdf.set_text_color(31, 78, 121)
                pdf.set_font("Helvetica", "U", 10)
                pdf.cell(0, 5, _latin1(f"Ubicación: {evento['lat']:.5f}, {evento['lon']:.5f}"),
                         link=maps_url(evento['lat'], evento['lon']),
                         new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.set_font("Helvetica", "", 11)
                pdf.set_text_color(0, 0, 0)
            if evento.get('imagen'):
                if pdf.get_y() > pdf.h - 60:
                    pdf.add_page()
                try:
                    pdf.image(BytesIO(evento['imagen']), w=70)
                    pdf.ln(1)
                except Exception:
                    pass
            pdf.ln(2)
    else:
        pdf.set_font("Helvetica", "", 11)
        pdf.multi_cell(0, 6, _latin1("Sin eventos registrados."),
                       new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    # Interacciones
    _seccion("Interacciones")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6, _latin1(interacciones or "Sin interacciones."),
                   new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    return bytes(pdf.output())

# =====================  Funciones de UI / persistencia  =====================
CLASIF_COLORS = {
    "CONFIDENCIAL": "#C0392B",
    "RESERVADO": "#E67E22",
    "USO INTERNO": "#2E86C1",
    "SIN CLASIFICAR": "#7F8C8D",
}


def badge_html(clasificacion):
    """Devuelve un badge HTML con color según el nivel de clasificación."""
    color = CLASIF_COLORS.get(clasificacion, "#7F8C8D")
    return (
        f'<span style="background:{color};color:#ffffff;padding:4px 12px;'
        f'border-radius:12px;font-size:0.85rem;font-weight:600;'
        f'letter-spacing:0.5px;">{clasificacion}</span>'
    )


def build_wa_summary(nombre_ao, fecha_operacion, zona, observaciones, eventos,
                     interacciones, max_obs=5000, max_desc=5000, max_items=40):
    """Genera un resumen de texto compacto para compartir por WhatsApp."""
    lines = [
        f"AO: {nombre_ao}",
        f"Fecha: {fecha_operacion.strftime('%d/%m/%Y')}",
        f"Zona: {zona}",
    ]
    if observaciones:
        obs = observaciones.strip()
        if len(obs) > max_obs:
            obs = obs[:max_obs].rsplit(' ', 1)[0] + "..."
        lines += ["", "Observaciones:", obs]
    if eventos:
        lines += ["", "Cronología:"]
        for e in sorted(eventos, key=hora_key)[:max_items]:
            hora_text = e.get('hora').strftime('%H:%M') if e.get('hora') else "--:--"
            desc = (e.get('descripcion') or "").strip()
            if len(desc) > max_desc:
                desc = desc[:max_desc].rsplit(' ', 1)[0] + "..."
            linea = f"{hora_text} - {desc}"
            if e.get('lat') is not None and e.get('lon') is not None:
                linea += f" 📍 {e['lat']:.5f},{e['lon']:.5f}"
            lines.append(linea)
        if len(eventos) > max_items:
            lines.append(f"...(+{len(eventos) - max_items} más)")
    if interacciones:
        inter = interacciones.strip()
        if len(inter) > max_obs:
            inter = inter[:max_obs].rsplit(' ', 1)[0] + "..."
        lines += ["", "Interacciones:", inter]
    return "\n".join(lines)


def serialize_report():
    """Serializa el informe actual (desde session_state) a una cadena JSON."""
    eventos_ser = []
    for e in sorted(st.session_state.get('eventos', []), key=hora_key):
        eventos_ser.append({
            'hora': e['hora'].strftime('%H:%M') if e.get('hora') else None,
            'descripcion': e.get('descripcion', ''),
            'lat': e.get('lat'),
            'lon': e.get('lon'),
            'imagen_b64': base64.b64encode(e['imagen']).decode('ascii') if e.get('imagen') else None,
        })
    fecha = st.session_state.get('fecha_operacion')
    data = {
        'version': 1,
        'nombre_ao': st.session_state.get('nombre_ao', ''),
        'fecha': fecha.isoformat() if fecha else None,
        'zona': st.session_state.get('zona', ''),
        'clasificacion': st.session_state.get('clasificacion', CLASIFICACIONES[0]),
        'observaciones': st.session_state.get('observaciones', ''),
        'interacciones': st.session_state.get('interacciones', ''),
        'eventos': eventos_ser,
    }
    return json.dumps(data, ensure_ascii=False, indent=2)


def cargar_callback():
    """Carga un informe desde el .json subido y rellena session_state."""
    archivo = st.session_state.get('uploader_json')
    if not archivo:
        st.session_state.load_msg = ("warning", "Selecciona un archivo .json primero.")
        return
    try:
        data = json.loads(archivo.getvalue().decode('utf-8'))
    except Exception as ex:
        st.session_state.load_msg = ("error", f"No se pudo leer el JSON: {ex}")
        return

    st.session_state.nombre_ao = data.get('nombre_ao', '')
    st.session_state.zona = data.get('zona', '')
    st.session_state.observaciones = data.get('observaciones', '')
    st.session_state.interacciones = data.get('interacciones', '')

    clasif = data.get('clasificacion', CLASIFICACIONES[0])
    st.session_state.clasificacion = clasif if clasif in CLASIFICACIONES else CLASIFICACIONES[0]

    fecha_str = data.get('fecha')
    if fecha_str:
        try:
            st.session_state.fecha_operacion = date.fromisoformat(fecha_str)
        except ValueError:
            pass

    eventos = []
    for e in data.get('eventos', []):
        hora = None
        if e.get('hora'):
            try:
                hora = datetime.strptime(e['hora'], '%H:%M').time()
            except ValueError:
                hora = None
        imagen = None
        if e.get('imagen_b64'):
            try:
                imagen = base64.b64decode(e['imagen_b64'])
            except Exception:
                imagen = None
        lat = e.get('lat')
        lon = e.get('lon')
        try:
            lat = float(lat) if lat is not None else None
            lon = float(lon) if lon is not None else None
        except (TypeError, ValueError):
            lat = lon = None
        eventos.append({'hora': hora, 'descripcion': e.get('descripcion', ''),
                        'lat': lat, 'lon': lon, 'imagen': imagen})
    st.session_state.eventos = eventos
    st.session_state.pop('editing_idx', None)
    st.session_state.load_msg = ("success", f"Informe cargado · {len(eventos)} evento(s).")


# =====================  Estado inicial  =====================
st.session_state.setdefault('eventos', [])
st.session_state.setdefault('upload_key', 0)
st.session_state.setdefault('hora_str', "00:00")
st.session_state.setdefault('descripcion', "")
st.session_state.setdefault('nombre_ao', "")
st.session_state.setdefault('zona', "")
st.session_state.setdefault('observaciones', "")
st.session_state.setdefault('interacciones', "")
st.session_state.setdefault('clasificacion', CLASIFICACIONES[0])
st.session_state.setdefault('fecha_operacion', datetime.now().date())
st.session_state.setdefault('docx_bytes', None)
st.session_state.setdefault('pdf_bytes', None)
st.session_state.setdefault('wa_summary', "")
st.session_state.setdefault('usar_ubicacion', False)
st.session_state.setdefault('lat_in', 0.0)
st.session_state.setdefault('lon_in', 0.0)
st.session_state.setdefault('map_pick_prev', None)

# Pequeños ajustes de estilo
st.markdown(
    """
    <style>
    .block-container {padding-top: 2.5rem;}
    div[data-testid="stMetricValue"] {font-size: 1.4rem;}
    </style>
    """,
    unsafe_allow_html=True,
)

# =====================  Barra lateral: datos + guardar/cargar  =====================
with st.sidebar:
    st.markdown("## 🕵️ Generador AO")
    st.caption("Informes de Análisis Operativo")
    st.divider()

    st.subheader("📌 Datos del informe")
    nombre_ao = st.text_input("Nombre (AO)", key="nombre_ao")
    fecha_operacion = st.date_input("Fecha", key="fecha_operacion")
    zona_operaciones = st.text_input("Zona", key="zona")
    clasificacion = st.selectbox("Clasificación", CLASIFICACIONES, key="clasificacion")

    st.divider()
    st.subheader("💾 Guardar / Cargar")
    st.download_button(
        "💾 Guardar informe (.json)",
        data=serialize_report(),
        file_name=f"AO_{nombre_ao or 'informe'}_{fecha_operacion.strftime('%Y%m%d')}.json",
        mime="application/json",
        use_container_width=True,
    )
    st.file_uploader("Cargar informe (.json)", type=['json'], key="uploader_json")
    st.button("📂 Cargar informe", on_click=cargar_callback, use_container_width=True)
    if st.session_state.get('load_msg'):
        nivel, texto = st.session_state.load_msg
        getattr(st, nivel)(texto)

# =====================  Cabecera + métricas  =====================
st.session_state.eventos.sort(key=hora_key)
n_eventos = len(st.session_state.eventos)
n_imagenes = sum(1 for e in st.session_state.eventos if e.get('imagen'))
n_geo = sum(1 for e in st.session_state.eventos
            if e.get('lat') is not None and e.get('lon') is not None)
horas = [e['hora'] for e in st.session_state.eventos if e.get('hora')]
franja = f"{min(horas).strftime('%H:%M')} – {max(horas).strftime('%H:%M')}" if horas else "—"

col_titulo, col_badge = st.columns([4, 1])
with col_titulo:
    st.title("Informe de Análisis Operativo")
    st.caption(
        f"{nombre_ao or '(sin nombre)'}  ·  {zona_operaciones or '(sin zona)'}"
        f"  ·  {fecha_es(fecha_operacion)}"
    )
with col_badge:
    st.markdown(badge_html(clasificacion), unsafe_allow_html=True)

m1, m2, m3, m4 = st.columns(4)
m1.metric("Eventos", n_eventos)
m2.metric("Con imagen", n_imagenes)
m3.metric("Geolocalizados", n_geo)
m4.metric("Franja horaria", franja)

st.divider()

tab_obs, tab_crono, tab_inter, tab_export = st.tabs(
    ["📋 Observaciones", "⏱️ Cronología", "👥 Interacciones", "📤 Exportar"]
)

# ---------- Observaciones ----------
with tab_obs:
    observaciones = st.text_area(
        "Observaciones principales",
        key="observaciones",
        height=220,
        placeholder="Describe las observaciones clave de la operación...",
    )

# ---------- Cronología ----------
with tab_crono:
    with st.expander("➕ Añadir nuevo evento", expanded=not st.session_state.eventos):
        col_h, col_d = st.columns([1, 3])
        with col_h:
            hora_str = st.text_input("Hora (HH:MM)", key="hora_str")
            try:
                hora = datetime.strptime(hora_str, "%H:%M").time()
            except ValueError:
                st.error("Formato HH:MM")
                hora = None
        with col_d:
            descripcion = st.text_area("Descripción del evento", height=100, key="descripcion")

        imagen = st.file_uploader(
            "Añadir imagen al evento (opcional)",
            type=['png', 'jpg', 'jpeg'],
            key=f"uploader_{st.session_state.upload_key}",
        )

        st.checkbox("📍 Añadir ubicación", key="usar_ubicacion")
        if st.session_state.usar_ubicacion:
            try:
                import folium
                from streamlit_folium import st_folium

                lat0 = st.session_state.lat_in or 40.4168
                lon0 = st.session_state.lon_in or -3.7038
                fmapa = folium.Map(location=[lat0, lon0], zoom_start=6)
                if st.session_state.lat_in or st.session_state.lon_in:
                    folium.Marker([st.session_state.lat_in, st.session_state.lon_in]).add_to(fmapa)
                ret_map = st_folium(
                    fmapa, height=300,
                    key=f"map_pick_{st.session_state.upload_key}",
                    returned_objects=["last_clicked"],
                )
                if ret_map and ret_map.get("last_clicked"):
                    punto = (round(ret_map["last_clicked"]["lat"], 6),
                             round(ret_map["last_clicked"]["lng"], 6))
                    if punto != st.session_state.get("map_pick_prev"):
                        st.session_state.map_pick_prev = punto
                        st.session_state.lat_in = punto[0]
                        st.session_state.lon_in = punto[1]
                        st.rerun()
                st.caption("Haz clic en el mapa para fijar el punto, o ajusta los valores abajo.")
            except ImportError:
                st.info("Para seleccionar en un mapa instala: pip install streamlit-folium. "
                        "Puedes introducir las coordenadas manualmente.")
            col_lat, col_lon = st.columns(2)
            with col_lat:
                st.number_input("Latitud", key="lat_in", format="%.6f", step=0.0001)
            with col_lon:
                st.number_input("Longitud", key="lon_in", format="%.6f", step=0.0001)

        def add_event_callback():
            try:
                h = datetime.strptime(st.session_state.hora_str, "%H:%M").time()
            except Exception:
                st.session_state.add_error = "Hora inválida. Introduce HH:MM antes de añadir el evento."
                return

            uploader_key = f"uploader_{st.session_state.upload_key}"
            imagen_widget = st.session_state.get(uploader_key)
            imagen_bytes = imagen_widget.getvalue() if imagen_widget else None

            if st.session_state.get('usar_ubicacion'):
                lat_val = st.session_state.get('lat_in')
                lon_val = st.session_state.get('lon_in')
            else:
                lat_val = lon_val = None

            st.session_state.eventos.append({
                'hora': h,
                'descripcion': st.session_state.descripcion,
                'lat': lat_val,
                'lon': lon_val,
                'imagen': imagen_bytes,
            })
            st.session_state.upload_key += 1
            st.session_state.hora_str = "00:00"
            st.session_state.descripcion = ""
            st.session_state.usar_ubicacion = False
            st.session_state.lat_in = 0.0
            st.session_state.lon_in = 0.0
            st.session_state.map_pick_prev = None
            st.session_state.add_error = ""
            st.session_state.add_success = "Evento añadido correctamente"

        st.button("➕ Añadir evento", on_click=add_event_callback)
        if st.session_state.get('add_error'):
            st.error(st.session_state.get('add_error'))
        if st.session_state.get('add_success'):
            st.success(st.session_state.get('add_success'))

    if st.session_state.eventos:
        for idx, evento in enumerate(st.session_state.eventos):
            col1, col2, col3 = st.columns([1, 3, 1])
            with col1:
                hora_display = evento['hora'].strftime('%H:%M') if evento.get('hora') else "—"
                st.write(f"**{hora_display}**")
            with col2:
                st.write(evento['descripcion'])
                if evento.get('lat') is not None and evento.get('lon') is not None:
                    st.markdown(
                        f"📍 [{evento['lat']:.5f}, {evento['lon']:.5f}]"
                        f"({maps_url(evento['lat'], evento['lon'])})"
                    )
                if evento['imagen']:
                    st.image(evento['imagen'], width=300)
            with col3:
                if st.button("✏️ Editar", key=f"edit_{idx}"):
                    st.session_state.editing_idx = idx
                if st.button("🗑️ Eliminar", key=f"delete_{idx}"):
                    st.session_state.eventos.pop(idx)
                    st.session_state.pop('editing_idx', None)
                    st.rerun()

            if 'editing_idx' in st.session_state and st.session_state.editing_idx == idx:
                with st.expander("Editar evento", expanded=True):
                    hora_edit = st.text_input(
                        "Hora (HH:MM)",
                        value=evento['hora'].strftime("%H:%M"),
                        key=f"hora_edit_{idx}",
                    )
                    desc_edit = st.text_area(
                        "Descripción",
                        value=evento['descripcion'],
                        key=f"desc_edit_{idx}",
                    )
                    imagen_edit = st.file_uploader(
                        "Actualizar imagen (opcional)",
                        type=['png', 'jpg', 'jpeg'],
                        key=f"img_edit_{idx}",
                    )

                    lat_e = lon_e = None
                    tiene_loc = st.checkbox(
                        "📍 Ubicación", value=evento.get('lat') is not None,
                        key=f"loc_chk_{idx}",
                    )
                    if tiene_loc:
                        cl1, cl2 = st.columns(2)
                        with cl1:
                            lat_e = st.number_input(
                                "Latitud", value=float(evento.get('lat') or 0.0),
                                format="%.6f", step=0.0001, key=f"lat_edit_{idx}",
                            )
                        with cl2:
                            lon_e = st.number_input(
                                "Longitud", value=float(evento.get('lon') or 0.0),
                                format="%.6f", step=0.0001, key=f"lon_edit_{idx}",
                            )

                    col_g, col_c = st.columns([1, 3])
                    with col_g:
                        if st.button("💾 Guardar", key=f"save_{idx}"):
                            try:
                                nueva_hora = datetime.strptime(hora_edit, "%H:%M").time()
                                st.session_state.eventos[idx]['hora'] = nueva_hora
                                st.session_state.eventos[idx]['descripcion'] = desc_edit
                                st.session_state.eventos[idx]['lat'] = lat_e if tiene_loc else None
                                st.session_state.eventos[idx]['lon'] = lon_e if tiene_loc else None
                                if imagen_edit:
                                    st.session_state.eventos[idx]['imagen'] = imagen_edit.getvalue()
                                del st.session_state.editing_idx
                                st.rerun()
                            except ValueError:
                                st.error("Formato de hora inválido (HH:MM)")
                    with col_c:
                        if st.button("❌ Cancelar", key=f"cancel_{idx}"):
                            del st.session_state.editing_idx
                            st.rerun()

        geo = [
            {"lat": e['lat'], "lon": e['lon']}
            for e in st.session_state.eventos
            if e.get('lat') is not None and e.get('lon') is not None
        ]
        if geo:
            import pandas as pd
            st.markdown("**🗺️ Mapa de situación**")
            st.map(pd.DataFrame(geo))

        if st.button("🧹 Limpiar cronología"):
            st.session_state.eventos = []
            st.session_state.pop('editing_idx', None)
            st.rerun()
    else:
        st.info("No hay eventos todavía. Añade el primero en el panel de arriba.")

# ---------- Interacciones ----------
with tab_inter:
    interacciones = st.text_area(
        "Interacciones relevantes",
        key="interacciones",
        height=220,
        placeholder="Describe las interacciones relevantes...",
    )

# ---------- Exportar ----------
with tab_export:
    nombre_archivo = f"AO_{nombre_ao or 'informe'}_{fecha_operacion.strftime('%Y%m%d')}"

    st.subheader("Documentos")
    col_word, col_pdf = st.columns(2)
    with col_word:
        if st.button("📝 Generar Word (DOCX)", use_container_width=True):
            st.session_state.docx_bytes = build_docx(
                nombre_ao, fecha_operacion, zona_operaciones, clasificacion,
                observaciones, st.session_state.eventos, interacciones,
            )
        if st.session_state.docx_bytes:
            st.download_button(
                "📥 Descargar DOCX",
                data=st.session_state.docx_bytes,
                file_name=f"{nombre_archivo}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
    with col_pdf:
        if st.button("📄 Generar PDF", use_container_width=True):
            try:
                st.session_state.pdf_bytes = build_pdf(
                    nombre_ao, fecha_operacion, zona_operaciones, clasificacion,
                    observaciones, st.session_state.eventos, interacciones,
                )
            except ImportError:
                st.session_state.pdf_bytes = None
                st.error("Instala fpdf2 para exportar a PDF: pip install fpdf2")
        if st.session_state.pdf_bytes:
            st.download_button(
                "📥 Descargar PDF",
                data=st.session_state.pdf_bytes,
                file_name=f"{nombre_archivo}.pdf",
                mime="application/pdf",
                use_container_width=True,
            )

    st.divider()
    st.subheader("Compartir")
    if st.button("✉️ Generar resumen WhatsApp"):
        st.session_state.wa_summary = build_wa_summary(
            nombre_ao, fecha_operacion, zona_operaciones, observaciones,
            st.session_state.eventos, interacciones,
        )
    if st.session_state.wa_summary:
        st.text_area(
            "Resumen para WhatsApp (copiar/pegar)",
            value=st.session_state.wa_summary,
            height=260,
            key="wa_show",
        )
        st.download_button(
            "📥 Descargar resumen (.txt)",
            data=st.session_state.wa_summary,
            file_name=f"{nombre_archivo}_resumen.txt",
            mime="text/plain",
        )
